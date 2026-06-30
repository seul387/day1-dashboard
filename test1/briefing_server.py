"""
저탄소에너지 브리핑 - 인터랙티브 서버
실행: python briefing_server.py
접속: http://localhost:5000
"""

from flask import Flask, jsonify, request, send_file
import urllib.request, urllib.parse, urllib.error
import json, re, io, os, sys
import anthropic
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

sys.stdout.reconfigure(encoding="utf-8")

app = Flask(__name__)

# -------------------------------------------------------
# API 키 및 데이터 경로
# -------------------------------------------------------
NEWSAPI_KEY       = "891d69f8cceb47759b5091fe99525333"
ANTHROPIC_API_KEY = os.environ.get("ANTHROPIC_API_KEY", "")

BASE_DIR      = os.path.dirname(os.path.abspath(__file__))
DATA_DIR      = os.path.join(BASE_DIR, "data")
ARCHIVE_DIR   = os.path.join(DATA_DIR, "archive")
BOOKMARKS_FILE = os.path.join(DATA_DIR, "bookmarks.json")

os.makedirs(ARCHIVE_DIR, exist_ok=True)
if not os.path.exists(BOOKMARKS_FILE):
    with open(BOOKMARKS_FILE, "w", encoding="utf-8") as f:
        json.dump([], f)

FROM_DATE = (datetime.now() - timedelta(days=2)).strftime("%Y-%m-%d")

KO_RSS_FEEDS = [
    "https://www.electimes.com/rss/allArticle.xml",
    "https://www.energy-news.co.kr/rss/allArticle.xml",
    "http://www.enewstoday.co.kr/rss/allArticle.xml",
    "https://www.gasnews.com/rss/allArticle.xml",
]

TOPICS = {
    "재생에너지":  {"icon": "☀️", "color": "#f59e0b", "rss_keywords": ["재생에너지","태양광","풍력","육상풍력","신재생","전기본","RE100","풍력발전","태양광발전"],   "newsapi_queries": ['"renewable energy" policy OR market OR investment OR "RE100" OR "solar" OR "onshore wind"'],   "exclude": []},
    "수소혼소":    {"icon": "🔥", "color": "#3b82f6", "rss_keywords": ["수소혼소","암모니아혼소","혼소발전"],   "newsapi_queries": ['"hydrogen co-firing" OR "hydrogen blending" OR "ammonia co-firing"'], "exclude": ["hydrogen bond","hydrogen transfer"]},
    "지열에너지":  {"icon": "🌋", "color": "#10b981", "rss_keywords": ["지열발전","지열에너지","지열"],          "newsapi_queries": ['"geothermal energy" OR "geothermal power" renewable'],        "exclude": ["earthquake","seismic","지진"]},
    "SMR":         {"icon": "⚛️", "color": "#8b5cf6", "rss_keywords": ["SMR","소형모듈원자로","소형원전"],       "newsapi_queries": ['"small modular reactor" OR "SMR" nuclear'],                 "exclude": []},
    "탄소배출권":  {"icon": "📊", "color": "#6b7280", "rss_keywords": ["탄소배출권","탄소시장","ETS","CBAM"],    "newsapi_queries": ['"carbon credit" OR "emission trading" OR "CBAM"'],         "exclude": []},
}

# -------------------------------------------------------
# 북마크 헬퍼
# -------------------------------------------------------
def load_bookmarks():
    try:
        with open(BOOKMARKS_FILE, encoding="utf-8") as f:
            return json.load(f)
    except:
        return []

def save_bookmarks(bms):
    with open(BOOKMARKS_FILE, "w", encoding="utf-8") as f:
        json.dump(bms, f, ensure_ascii=False, indent=2)

# -------------------------------------------------------
# 아카이브 헬퍼
# -------------------------------------------------------
def save_archive(date, data):
    path = os.path.join(ARCHIVE_DIR, f"{date}.json")
    if not os.path.exists(path):
        with open(path, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        print(f"아카이브 저장: {date}.json")

def list_archive_dates():
    files = [f.replace(".json", "") for f in os.listdir(ARCHIVE_DIR) if f.endswith(".json")]
    return sorted(files, reverse=True)

def load_archive(date):
    path = os.path.join(ARCHIVE_DIR, f"{date}.json")
    if not os.path.exists(path):
        return None
    with open(path, encoding="utf-8") as f:
        return json.load(f)

# -------------------------------------------------------
# 뉴스 수집
# -------------------------------------------------------
_rss_cache = []

def load_rss():
    global _rss_cache
    if _rss_cache:
        return
    SOURCE_MAP = {"electimes": "전기신문", "energy-news": "에너지신문",
                  "enewstoday": "이뉴스투데이", "gasnews": "가스신문"}
    for feed_url in KO_RSS_FEEDS:
        try:
            req = urllib.request.Request(feed_url, headers={"User-Agent": "Mozilla/5.0"})
            with urllib.request.urlopen(req, timeout=10) as r:
                raw = r.read().decode("utf-8", errors="ignore")
            for item in re.findall(r"<item>(.*?)</item>", raw, re.DOTALL):
                def tag(t):
                    m = re.search(rf"<{t}[^>]*>(?:<!\[CDATA\[)?(.*?)(?:\]\]>)?</{t}>", item, re.DOTALL)
                    return m.group(1).strip() if m else ""
                title = tag("title")
                link  = (tag("link") or tag("guid")).strip()
                desc  = re.sub(r"<[^>]+>", "", tag("description"))[:150]
                pub   = tag("pubDate")[:16]
                src   = next((v for k, v in SOURCE_MAP.items() if k in feed_url), "국내언론")
                if title and link:
                    _rss_cache.append({"title": title, "description": desc, "url": link,
                                       "publishedAt": pub, "source": {"name": src}, "lang": "ko"})
        except Exception as e:
            print(f"RSS 오류: {e}")

def fetch_newsapi(query, page_size=2, from_date=None, to_date=None):
    params = {"q": query, "from": from_date or FROM_DATE, "sortBy": "publishedAt",
              "pageSize": str(page_size), "language": "en", "apiKey": NEWSAPI_KEY}
    if to_date:
        params["to"] = to_date
    url = "https://newsapi.org/v2/everything?" + urllib.parse.urlencode(params)
    try:
        req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
        with urllib.request.urlopen(req, timeout=10) as r:
            data = json.loads(r.read().decode())
            for a in data.get("articles", []):
                a["lang"] = "en"
            return data.get("articles", [])
    except:
        return []

def collect_all(from_date=None, to_date=None):
    load_rss()
    results = {}
    for topic, cfg in TOPICS.items():
        seen, articles = set(), []
        # RSS: 날짜 필터 적용 (to_date 기준)
        for a in _rss_cache:
            if to_date:
                pub = (a.get("publishedAt") or "")[:10]
                if pub and pub > to_date:
                    continue
            text = (a["title"] + " " + a["description"]).lower()
            if any(k.lower() in text for k in cfg["rss_keywords"]) and a["url"] not in seen:
                if not any(e.lower() in text for e in cfg["exclude"]):
                    seen.add(a["url"]); articles.append(a)
        for q in cfg["newsapi_queries"]:
            for a in fetch_newsapi(q, from_date=from_date, to_date=to_date):
                text = (a.get("title","") + " " + a.get("description","")).lower()
                if a["url"] not in seen and not any(e.lower() in text for e in cfg["exclude"]):
                    seen.add(a["url"]); articles.append(a)
        results[topic] = articles[:5]
    return results

# -------------------------------------------------------
# 90일 이전 아카이브 자동 정리
# -------------------------------------------------------
def cleanup_old_archives(keep_days=90):
    cutoff = (datetime.now() - timedelta(days=keep_days)).strftime("%Y-%m-%d")
    removed = 0
    for fname in os.listdir(ARCHIVE_DIR):
        if not fname.endswith(".json"):
            continue
        date_str = fname.replace(".json", "")
        if date_str < cutoff:
            os.remove(os.path.join(ARCHIVE_DIR, fname))
            removed += 1
    if removed:
        print(f"오래된 아카이브 {removed}건 삭제 완료 ({keep_days}일 이전)")
    else:
        print(f"정리할 오래된 아카이브 없음 ({keep_days}일 기준)")

# -------------------------------------------------------
# 누락 아카이브 자동 보완 (최근 7일 체크)
# -------------------------------------------------------
def backfill_missing_archives():
    today = datetime.now()
    missing = []
    for i in range(1, 8):
        target = (today - timedelta(days=i)).strftime("%Y-%m-%d")
        if not os.path.exists(os.path.join(ARCHIVE_DIR, f"{target}.json")):
            missing.append(target)

    if not missing:
        print("누락 아카이브 없음 (최근 7일 모두 정상)")
        return

    for target in missing:
        print(f"누락 아카이브 보완 중: {target} ...")
        try:
            data = collect_all(from_date=target, to_date=target)
            total = sum(len(v) for v in data.values())
            path = os.path.join(ARCHIVE_DIR, f"{target}.json")
            with open(path, "w", encoding="utf-8") as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
            print(f"  → {target}: {total}건 저장 완료")
        except Exception as e:
            print(f"  → {target} 보완 실패: {e}")

# -------------------------------------------------------
# Claude 호출
# -------------------------------------------------------
def call_claude(prompt):
    client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
    message = client.messages.create(
        model="claude-opus-4-8",
        max_tokens=2048,
        messages=[{"role": "user", "content": prompt}]
    )
    return message.content[0].text

# -------------------------------------------------------
# Word 문서 생성
# -------------------------------------------------------
def build_docx(draft_text, today):
    doc = Document()
    title = doc.add_heading("저탄소에너지 브리핑", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(0x1e, 0x3a, 0x5f)
    sub = doc.add_paragraph(f"작성일: {today}")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(10)
    sub.runs[0].font.color.rgb = RGBColor(0x71, 0x80, 0x96)
    doc.add_paragraph()
    for line in draft_text.split("\n"):
        line = line.strip()
        if not line:
            doc.add_paragraph(); continue
        if line.startswith("##"):
            h = doc.add_heading(line.lstrip("#").strip(), level=2)
            h.runs[0].font.color.rgb = RGBColor(0x2d, 0x6a, 0x4f)
        elif line.startswith("#"):
            h = doc.add_heading(line.lstrip("#").strip(), level=1)
            h.runs[0].font.color.rgb = RGBColor(0x1e, 0x3a, 0x5f)
        else:
            doc.add_paragraph(line)
    buf = io.BytesIO()
    doc.save(buf); buf.seek(0)
    return buf

# -------------------------------------------------------
# 공통 CSS / 네비게이션
# -------------------------------------------------------
COMMON_CSS = """
    *{box-sizing:border-box;margin:0;padding:0}
    body{font-family:'Segoe UI','Malgun Gothic',sans-serif;background:#f0f4f8;color:#1a202c}
    header{background:linear-gradient(135deg,#1e3a5f,#2d6a4f);color:white;padding:24px 40px}
    header h1{font-size:1.6rem;font-weight:700}
    header p{margin-top:4px;opacity:.8;font-size:.9rem}
    nav.page-nav{background:#1a2f4a;display:flex;gap:4px;padding:0 36px}
    nav.page-nav a{color:rgba(255,255,255,.65);text-decoration:none;padding:10px 18px;font-size:.88rem;
                   font-weight:600;border-bottom:3px solid transparent;transition:all .2s}
    nav.page-nav a:hover{color:white}
    nav.page-nav a.active{color:white;border-bottom-color:#4ade80}
    .toolbar{background:white;border-bottom:1px solid #e2e8f0;padding:12px 40px;display:flex;
             gap:12px;align-items:center;position:sticky;top:0;z-index:100;box-shadow:0 2px 8px rgba(0,0,0,.06)}
    .btn{padding:8px 18px;border-radius:6px;border:none;cursor:pointer;font-size:.9rem;font-weight:600;transition:all .2s}
    .btn-primary{background:#2d6a4f;color:white} .btn-primary:hover{background:#1e3a5f}
    .btn-secondary{background:#e2e8f0;color:#4a5568} .btn-secondary:hover{background:#cbd5e0}
    .btn-outline{background:white;color:#2d6a4f;border:2px solid #2d6a4f} .btn-outline:hover{background:#f0fff4}
    .btn-danger{background:#fff5f5;color:#e53e3e;border:1px solid #fed7d7} .btn-danger:hover{background:#fed7d7}
    .badge{background:#e53e3e;color:white;border-radius:99px;padding:2px 8px;font-size:.75rem;margin-left:6px}
    main{max-width:1000px;margin:28px auto;padding:0 20px 80px}
    .topic-section{margin-bottom:36px}
    .topic-title{font-size:1.1rem;font-weight:700;margin-bottom:14px;padding-left:12px;border-left:4px solid #ccc}
    .cards{display:grid;grid-template-columns:repeat(auto-fill,minmax(290px,1fr));gap:14px}
    .card{background:white;border-radius:10px;padding:16px;box-shadow:0 1px 4px rgba(0,0,0,.08);
          border:2px solid transparent;transition:all .2s}
    .card.selected{border-color:#2d6a4f;background:#f0fff4}
    .card.bookmarked{border-color:#f59e0b}
    .card-header{display:flex;align-items:flex-start;gap:8px;margin-bottom:10px}
    .card-check{width:20px;height:20px;accent-color:#2d6a4f;cursor:pointer;flex-shrink:0;margin-top:2px}
    .card-meta{font-size:.72rem;color:#718096}
    .bm-btn{background:none;border:none;cursor:pointer;font-size:1.1rem;padding:0 2px;
            line-height:1;flex-shrink:0;opacity:.4;transition:opacity .2s}
    .bm-btn:hover{opacity:1}
    .bm-btn.on{opacity:1}
    .card-title{font-size:.9rem;font-weight:600;color:#2b6cb0;line-height:1.4;margin-bottom:6px}
    .card-title a{color:inherit;text-decoration:none} .card-title a:hover{text-decoration:underline}
    .card-desc{font-size:.8rem;color:#4a5568;line-height:1.5;margin-bottom:10px}
    .card-memo{width:100%;padding:6px 8px;border:1px solid #e2e8f0;border-radius:6px;
               font-size:.8rem;resize:vertical;min-height:50px;font-family:inherit;color:#4a5568}
    .card-memo:focus{outline:none;border-color:#2d6a4f}
    .empty{color:#a0aec0;font-size:.9rem;padding:20px 0;text-align:center}
    .draft-panel{position:fixed;right:0;top:0;width:420px;height:100vh;background:white;
                 box-shadow:-4px 0 20px rgba(0,0,0,.15);transform:translateX(100%);
                 transition:transform .3s;z-index:200;display:flex;flex-direction:column}
    .draft-panel.open{transform:translateX(0)}
    .draft-header{background:#1e3a5f;color:white;padding:20px;display:flex;justify-content:space-between;align-items:center}
    .draft-header h3{font-size:1rem}
    .draft-close{background:none;border:none;color:white;font-size:1.3rem;cursor:pointer}
    .draft-body{flex:1;overflow-y:auto;padding:20px}
    #draft-textarea{width:100%;height:100%;min-height:400px;border:1px solid #e2e8f0;border-radius:8px;
                    padding:14px;font-size:.85rem;line-height:1.7;font-family:inherit;resize:none}
    #draft-textarea:focus{outline:none;border-color:#2d6a4f}
    .draft-footer{padding:16px 20px;border-top:1px solid #e2e8f0;display:flex;gap:10px}
    .loading{display:none;text-align:center;padding:40px;color:#718096}
    .spinner{width:32px;height:32px;border:3px solid #e2e8f0;border-top-color:#2d6a4f;
             border-radius:50%;animation:spin 1s linear infinite;margin:0 auto 12px}
    @keyframes spin{to{transform:rotate(360deg)}}
    .archive-list{display:flex;flex-direction:column;gap:10px;max-width:600px}
    .archive-item{background:white;border-radius:10px;padding:16px 20px;box-shadow:0 1px 4px rgba(0,0,0,.08);
                  display:flex;align-items:center;justify-content:space-between;cursor:pointer;
                  border:2px solid transparent;transition:all .2s}
    .archive-item:hover{border-color:#2d6a4f;background:#f0fff4}
    .archive-date{font-weight:700;font-size:1rem;color:#1e3a5f}
    .archive-count{font-size:.8rem;color:#718096}
    footer{text-align:center;font-size:.78rem;color:#a0aec0;padding:20px}
"""

def nav_html(active):
    links = [("/", "📰 오늘 브리핑"), ("/bookmarks", "⭐ 북마크"), ("/archive", "📅 아카이브")]
    items = "".join(
        f'<a href="{href}" class="nav-link{"  active" if href == active else ""}">{label}</a>'
        for href, label in links
    )
    return f'<nav class="page-nav">{items}</nav>'

# -------------------------------------------------------
# 뉴스 데이터 (서버 시작 시 1회 수집)
# -------------------------------------------------------
print("뉴스 수집 중...")
NEWS_DATA = collect_all()
total = sum(len(v) for v in NEWS_DATA.values())
print(f"수집 완료: 총 {total}건")

TODAY = datetime.now().strftime("%Y-%m-%d")
save_archive(TODAY, NEWS_DATA)

cleanup_old_archives(keep_days=90)
backfill_missing_archives()

# -------------------------------------------------------
# Flask 라우트 - 메인
# -------------------------------------------------------
@app.route("/")
def index():
    today = TODAY
    topics_json     = json.dumps(NEWS_DATA, ensure_ascii=False)
    topics_cfg_json = json.dumps(
        {k: {"icon": v["icon"], "color": v["color"]} for k, v in TOPICS.items()},
        ensure_ascii=False
    )
    return f"""<!DOCTYPE html>
<html lang="ko">
<head>
  <meta charset="UTF-8">
  <meta name="viewport" content="width=device-width,initial-scale=1.0">
  <title>저탄소에너지 브리핑 [{today}]</title>
  <style>{COMMON_CSS}</style>
</head>
<body>

<header>
  <h1>저탄소에너지 브리핑</h1>
  <p>{today} · 재생에너지 · 수소혼소 · 지열에너지 · SMR · 탄소배출권</p>
</header>
{nav_html("/")}

<div class="toolbar">
  <button class="btn btn-secondary" onclick="selectAll()">전체 선택</button>
  <button class="btn btn-secondary" onclick="clearAll()">선택 해제</button>
  <span style="color:#718096;font-size:.85rem">선택된 기사: <strong id="count">0</strong>건</span>
  <div style="flex:1"></div>
</div>

<main id="main-content"></main>

<footer>자동 수집 브리핑 · RSS + NewsAPI + Claude 분석</footer>

<script>
const NEWS = {topics_json};
const CFG  = {topics_cfg_json};
let bookmarkedUrls = new Set();

async function loadBookmarkState() {{
  const res = await fetch('/api/bookmarks');
  const bms = await res.json();
  bookmarkedUrls = new Set(bms.map(b => b.url));
  document.querySelectorAll('.bm-btn').forEach(btn => {{
    if (bookmarkedUrls.has(btn.dataset.url)) btn.classList.add('on');
  }});
}}

function renderCards() {{
  const main = document.getElementById('main-content');
  main.innerHTML = '';
  for (const [topic, articles] of Object.entries(NEWS)) {{
    const cfg = CFG[topic] || {{}};
    const section = document.createElement('section');
    section.className = 'topic-section';
    section.innerHTML = `<h2 class="topic-title" style="border-left-color:${{cfg.color||'#ccc'}}">${{cfg.icon||''}} ${{topic}}</h2>`;
    const grid = document.createElement('div');
    grid.className = 'cards';
    if (!articles.length) {{
      grid.innerHTML = '<p class="empty">수집된 기사 없음</p>';
    }} else {{
      articles.forEach((a, i) => {{
        const flag = a.lang === 'ko' ? '🇰🇷' : '🌐';
        const pub  = (a.publishedAt || '').slice(0, 10);
        const src  = (a.source || {{}}).name || '';
        const card = document.createElement('div');
        card.className = 'card';
        card.dataset.topic = topic;
        card.dataset.idx   = i;
        card.innerHTML = `
          <div class="card-header">
            <input type="checkbox" class="card-check" id="chk-${{topic}}-${{i}}">
            <div style="flex:1">
              <div class="card-meta">${{flag}} ${{src}} · ${{pub}}</div>
            </div>
            <button class="bm-btn" data-url="${{a.url}}" title="북마크">⭐</button>
          </div>
          <div class="card-title"><a href="${{a.url}}" target="_blank">${{a.title||'(제목 없음)'}}</a></div>
          <p class="card-desc">${{(a.description||'').slice(0,100)}}</p>
          <textarea class="card-memo" placeholder="메모 입력 (선택사항)..." rows="2"></textarea>`;
        card.querySelector('.card-check').addEventListener('change', function() {{
          card.classList.toggle('selected', this.checked);
          updateCount();
        }});
        card.querySelector('.bm-btn').addEventListener('click', async function(e) {{
          e.stopPropagation();
          const url = a.url;
          const memo = card.querySelector('.card-memo').value.trim();
          if (bookmarkedUrls.has(url)) {{
            await fetch('/api/bookmark/remove', {{method:'POST',headers:{{'Content-Type':'application/json'}},body:JSON.stringify({{url}})}});
            bookmarkedUrls.delete(url);
            this.classList.remove('on');
            card.classList.remove('bookmarked');
          }} else {{
            await fetch('/api/bookmark/add', {{method:'POST',headers:{{'Content-Type':'application/json'}},
              body:JSON.stringify({{topic, article:a, memo}})}});
            bookmarkedUrls.add(url);
            this.classList.add('on');
            card.classList.add('bookmarked');
          }}
        }});
        grid.appendChild(card);
      }});
    }}
    section.appendChild(grid);
    main.appendChild(section);
  }}
}}

function updateCount() {{
  const n = document.querySelectorAll('.card-check:checked').length;
  document.getElementById('count').textContent = n;
  document.getElementById('sel-badge').textContent = n;
}}

function selectAll() {{
  document.querySelectorAll('.card-check').forEach(c => {{ c.checked = true; c.dispatchEvent(new Event('change')); }});
}}
function clearAll() {{
  document.querySelectorAll('.card-check').forEach(c => {{ c.checked = false; c.dispatchEvent(new Event('change')); }});
}}

renderCards();
loadBookmarkState();
</script>
</body>
</html>"""


# -------------------------------------------------------
# Flask 라우트 - 북마크 페이지
# -------------------------------------------------------
@app.route("/bookmarks")
def bookmarks_page():
    bms = load_bookmarks()
    bms_json = json.dumps(bms, ensure_ascii=False)
    count = len(bms)
    topics_cfg_json = json.dumps(
        {k: {"icon": v["icon"], "color": v["color"]} for k, v in TOPICS.items()},
        ensure_ascii=False
    )
    return f"""<!DOCTYPE html>
<html lang="ko">
<head>
  <meta charset="UTF-8">
  <meta name="viewport" content="width=device-width,initial-scale=1.0">
  <title>북마크 · 저탄소에너지 브리핑</title>
  <style>{COMMON_CSS}</style>
</head>
<body>

<header>
  <h1>저탄소에너지 브리핑</h1>
  <p>⭐ 북마크 · 저장된 기사 {count}건</p>
</header>
{nav_html("/bookmarks")}

<main>
  <div id="bm-content"></div>
</main>

<footer>자동 수집 브리핑 · RSS + NewsAPI + Claude 분석</footer>

<script>
const BMS = {bms_json};
const CFG = {topics_cfg_json};

function renderBookmarks() {{
  const el = document.getElementById('bm-content');
  if (!BMS.length) {{
    el.innerHTML = '<p class="empty" style="margin-top:60px">저장된 북마크가 없습니다.<br><a href="/" style="color:#2d6a4f">오늘 브리핑</a>에서 ⭐를 눌러 저장하세요.</p>';
    return;
  }}
  const grouped = {{}};
  BMS.forEach(b => {{ (grouped[b.topic] = grouped[b.topic]||[]).push(b); }});
  let html = '';
  for (const [topic, arts] of Object.entries(grouped)) {{
    const cfg = CFG[topic] || {{}};
    html += `<section class="topic-section">
      <h2 class="topic-title" style="border-left-color:${{cfg.color||'#ccc'}}">${{cfg.icon||''}} ${{topic}}</h2>
      <div class="cards">`;
    arts.forEach((a, i) => {{
      const flag = a.lang === 'ko' ? '🇰🇷' : '🌐';
      const src  = (a.source||{{}}).name||'';
      const pub  = (a.publishedAt||'').slice(0,10);
      html += `<div class="card bookmarked" id="bm-${{i}}-${{topic}}">
        <div class="card-header">
          <div style="flex:1"><div class="card-meta">${{flag}} ${{src}} · ${{pub}}</div></div>
          <div class="card-meta" style="color:#b7791f">저장: ${{a.saved_at||''}}</div>
        </div>
        <div class="card-title"><a href="${{a.url}}" target="_blank">${{a.title||'(제목 없음)'}}</a></div>
        <p class="card-desc">${{(a.description||'').slice(0,120)}}</p>
        ${{a.memo ? `<p style="font-size:.78rem;color:#b7791f;margin-top:6px">📝 ${{a.memo}}</p>` : ''}}
        <button class="btn btn-danger" style="margin-top:10px;width:100%;font-size:.8rem"
          onclick="removeBookmark('${{encodeURIComponent(a.url)}}', this)">북마크 해제</button>
      </div>`;
    }});
    html += '</div></section>';
  }}
  el.innerHTML = html;
}}

async function removeBookmark(encodedUrl, btn) {{
  const url = decodeURIComponent(encodedUrl);
  await fetch('/api/bookmark/remove', {{method:'POST',headers:{{'Content-Type':'application/json'}},body:JSON.stringify({{url}})}});
  btn.closest('.card').style.opacity = '0.3';
  setTimeout(() => location.reload(), 400);
}}

renderBookmarks();
</script>
</body>
</html>"""


# -------------------------------------------------------
# Flask 라우트 - 아카이브 목록
# -------------------------------------------------------
@app.route("/archive")
def archive_page():
    dates = list_archive_dates()
    archive_counts = {}
    for d in dates:
        data = load_archive(d)
        if data:
            archive_counts[d] = sum(len(v) for v in data.values())

    rows_html = ""
    for d in dates:
        cnt = archive_counts.get(d, 0)
        label = "오늘" if d == TODAY else ""
        rows_html += f"""<div class="archive-item" onclick="location.href='/archive/{d}'">
          <div>
            <div class="archive-date">{d} {"<span style='background:#2d6a4f;color:white;padding:2px 8px;border-radius:99px;font-size:.72rem;margin-left:8px'>오늘</span>" if label else ""}</div>
            <div class="archive-count" style="margin-top:4px">기사 {cnt}건</div>
          </div>
          <span style="color:#a0aec0;font-size:1.2rem">›</span>
        </div>"""

    if not rows_html:
        rows_html = '<p class="empty">아카이브가 없습니다. 서버를 매일 실행하면 자동으로 쌓입니다.</p>'

    return f"""<!DOCTYPE html>
<html lang="ko">
<head>
  <meta charset="UTF-8">
  <meta name="viewport" content="width=device-width,initial-scale=1.0">
  <title>아카이브 · 저탄소에너지 브리핑</title>
  <style>{COMMON_CSS}</style>
</head>
<body>

<header>
  <h1>저탄소에너지 브리핑</h1>
  <p>📅 아카이브 · {len(dates)}일치 보관 중</p>
</header>
{nav_html("/archive")}

<main>
  <div class="archive-list">
    {rows_html}
  </div>
</main>

<footer>자동 수집 브리핑 · RSS + NewsAPI + Claude 분석</footer>
</body>
</html>"""


# -------------------------------------------------------
# Flask 라우트 - 아카이브 날짜별
# -------------------------------------------------------
@app.route("/archive/<date>")
def archive_date_page(date):
    data = load_archive(date)
    if not data:
        return f"<h2>아카이브 없음: {date}</h2><a href='/archive'>돌아가기</a>", 404

    topics_cfg_json = json.dumps(
        {k: {"icon": v["icon"], "color": v["color"]} for k, v in TOPICS.items()},
        ensure_ascii=False
    )
    data_json = json.dumps(data, ensure_ascii=False)
    total = sum(len(v) for v in data.values())

    return f"""<!DOCTYPE html>
<html lang="ko">
<head>
  <meta charset="UTF-8">
  <meta name="viewport" content="width=device-width,initial-scale=1.0">
  <title>{date} 브리핑 · 아카이브</title>
  <style>{COMMON_CSS}</style>
</head>
<body>

<header>
  <h1>저탄소에너지 브리핑</h1>
  <p>📅 {date} 아카이브 · 기사 {total}건</p>
</header>
{nav_html("/archive")}

<div class="toolbar">
  <button class="btn btn-secondary" onclick="history.back()">← 목록으로</button>
  <span style="color:#718096;font-size:.85rem;margin-left:8px">{date} 수집 기사</span>
</div>

<main id="main-content"></main>

<footer>자동 수집 브리핑 · RSS + NewsAPI + Claude 분석</footer>

<script>
const NEWS = {data_json};
const CFG  = {topics_cfg_json};

function renderCards() {{
  const main = document.getElementById('main-content');
  main.innerHTML = '';
  for (const [topic, articles] of Object.entries(NEWS)) {{
    const cfg = CFG[topic] || {{}};
    const section = document.createElement('section');
    section.className = 'topic-section';
    section.innerHTML = `<h2 class="topic-title" style="border-left-color:${{cfg.color||'#ccc'}}">${{cfg.icon||''}} ${{topic}}</h2>`;
    const grid = document.createElement('div');
    grid.className = 'cards';
    if (!articles.length) {{
      grid.innerHTML = '<p class="empty">수집된 기사 없음</p>';
    }} else {{
      articles.forEach(a => {{
        const flag = a.lang === 'ko' ? '🇰🇷' : '🌐';
        const pub  = (a.publishedAt||'').slice(0,10);
        const src  = (a.source||{{}}).name||'';
        const card = document.createElement('div');
        card.className = 'card';
        card.innerHTML = `
          <div class="card-header">
            <div style="flex:1"><div class="card-meta">${{flag}} ${{src}} · ${{pub}}</div></div>
          </div>
          <div class="card-title"><a href="${{a.url}}" target="_blank">${{a.title||'(제목 없음)'}}</a></div>
          <p class="card-desc">${{(a.description||'').slice(0,120)}}</p>`;
        grid.appendChild(card);
      }});
    }}
    section.appendChild(grid); main.appendChild(section);
  }}
}}
renderCards();
</script>
</body>
</html>"""


# -------------------------------------------------------
# API - 북마크
# -------------------------------------------------------
@app.route("/api/bookmarks")
def api_bookmarks():
    return jsonify(load_bookmarks())

@app.route("/api/bookmark/add", methods=["POST"])
def api_bookmark_add():
    body = request.get_json()
    article = body.get("article", {})
    topic   = body.get("topic", "")
    memo    = body.get("memo", "")
    if not article.get("url"):
        return jsonify({"ok": False})
    bms = load_bookmarks()
    if any(b["url"] == article["url"] for b in bms):
        return jsonify({"ok": True, "action": "exists"})
    entry = {**article, "topic": topic, "memo": memo,
             "saved_at": datetime.now().strftime("%Y-%m-%d %H:%M")}
    bms.append(entry)
    save_bookmarks(bms)
    return jsonify({"ok": True, "action": "added"})

@app.route("/api/bookmark/remove", methods=["POST"])
def api_bookmark_remove():
    body = request.get_json()
    url  = body.get("url", "")
    bms  = [b for b in load_bookmarks() if b["url"] != url]
    save_bookmarks(bms)
    return jsonify({"ok": True})


# -------------------------------------------------------
# API - 아카이브
# -------------------------------------------------------
@app.route("/api/archive")
def api_archive():
    return jsonify(list_archive_dates())

@app.route("/api/archive/<date>")
def api_archive_date(date):
    data = load_archive(date)
    if not data:
        return jsonify({"error": "없음"}), 404
    return jsonify(data)


# -------------------------------------------------------
# API - 초안 생성 / Word 다운로드
# -------------------------------------------------------
@app.route("/api/generate-draft", methods=["POST"])
def generate_draft():
    body = request.get_json()
    articles = body.get("articles", [])
    date = body.get("date", datetime.now().strftime("%Y-%m-%d"))
    if not articles:
        return jsonify({"error": "선택된 기사 없음"})

    lines = []
    for a in articles:
        memo_part = f"\n  [담당자 메모: {a['memo']}]" if a.get("memo") else ""
        lines.append(f"[{a['topic']}] {a['title']}\n  요약: {a.get('description','')[:100]}{memo_part}")
    article_block = "\n\n".join(lines)

    prompt = f"""당신은 에너지 전문 브리핑 작성자입니다.
아래 수집된 기사들을 바탕으로 {date} 저탄소에너지 브리핑 초안을 작성하세요.

[수집 기사]
{article_block}

[작성 규칙]
- 분야별로 섹션을 나누어 작성 (## 분야명 형식)
- 각 섹션에 2~4문장의 동향 설명
- 객관적이고 전문적인 보고서 어투 사용 (~됨, ~로 전망됨, ~발표됨)
- 담당자 메모가 있으면 해당 내용을 분석에 반영
- 마지막에 ## 종합 시사점 섹션 추가 (3~4문장)
- 불필요한 인사말 없이 바로 본문 시작

지금 바로 초안을 작성하세요:"""

    try:
        draft = call_claude(prompt)
        return jsonify({"draft": draft})
    except Exception as e:
        draft_lines = [f"# 저탄소에너지 브리핑  [{date}]\n"]
        grouped = {}
        for a in articles:
            grouped.setdefault(a["topic"], []).append(a)
        for topic, arts in grouped.items():
            draft_lines.append(f"\n## {topic}\n")
            for a in arts:
                memo = f" ({a['memo']})" if a.get("memo") else ""
                draft_lines.append(f"- {a['title']}{memo}")
        draft_lines.append("\n## 종합 시사점\n")
        draft_lines.append(f"※ Claude API 연결 실패 ({e})로 템플릿 초안이 생성되었습니다.")
        return jsonify({"draft": "\n".join(draft_lines)})


@app.route("/api/download-docx", methods=["POST"])
def download_docx():
    body  = request.get_json()
    draft = body.get("draft", "")
    date  = body.get("date", datetime.now().strftime("%Y-%m-%d"))
    buf   = build_docx(draft, date)
    return send_file(buf, as_attachment=True,
                     download_name=f"저탄소에너지브리핑_{date}.docx",
                     mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


if __name__ == "__main__":
    print("브라우저에서 http://localhost:5000 접속하세요")
    app.run(debug=False, port=5000)
